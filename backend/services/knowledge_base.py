"""
Knowledge Base Service for Catalyst
Handles document management, processing, and semantic search integration
"""

import asyncio
import logging
import hashlib
import json
import os
import mimetypes
try:
    from typing import List, Dict, Any, Optional, Union, Tuple, IO, Callable
    from datetime import datetime, timezone
    from dataclasses import dataclass, asdict
    from enum import Enum
    from pathlib import Path
    import re
    import functools
except ImportError:
    pass


except ImportError:
    pass

    pass

    pass
# Import the vector search service
try:
    from .vector_search import VectorSearchService, SearchResult


except ImportError:
    pass

    pass

    pass
# Try importing document processing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentType(str, Enum):
    """Document types for categorization"""
    GUIDANCE = "guidance"
    INSTRUCTION = "instruction"
    CONTEXT = "context"
    REFERENCE = "reference"
    CASE_STUDY = "case_study"
    TEMPLATE = "template"
    RESEARCH = "research"
    MANUAL = "manual"
    FAQ = "faq"
    OTHER = "other"

class ProcessingStatus(str, Enum):
    """Document processing status"""
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"
    INDEXED = "indexed"

def handle_kb_errors(func):
    """
    Decorator for handling knowledge base service errors
    
    Args:
        func: The async function to wrap
        
    Returns:
        Wrapped function with error handling
    """
    @functools.wraps(func)
    async def wrapper(*args, **kwargs):
        try:
            return await func(*args, **kwargs)
        except Exception as e:
            # Get the function name for better error context
            func_name = func.__name__
            
            # Log the error with detailed context
            logger.error(f"Knowledge base error in {func_name}: {str(e)}", 
                         extra={
                             "args": args[1:],  # Skip self
                             "kwargs": kwargs,
                             "error_type": type(e).__name__
                         },
                         exc_info=True)
            
            # Return a proper error response instead of raising
            if func_name.startswith("search"):
                return []
            elif func_name.startswith("get_"):
                return None
            elif func_name.startswith("index_") or func_name.startswith("process_"):
                return {
                    "success": False,
                    "error": str(e),
                    "error_type": type(e).__name__
                }
            else:
                # For update/delete operations
                return {
                    "success": False,
                    "error": str(e)
                }
    
    return wrapper

@dataclass
class DocumentChunk:
    """A chunk of a larger document"""
    chunk_id: str
    content: str
    chunk_index: int
    total_chunks: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any]

@dataclass
class ProcessingResult:
    """Result of document processing"""
    document_id: str
    success: bool
    chunks_created: int
    processing_time: float
    errors: List[str]
    metadata: Dict[str, Any]

@dataclass
class KnowledgeDocument:
    """Knowledge base document"""
    id: str
    title: str
    content: str
    document_type: DocumentType
    file_path: Optional[str] = None
    file_type: Optional[str] = None
    file_size: Optional[int] = None
    tags: Optional[List[str]] = None
    metadata: Optional[Dict[str, Any]] = None
    created_at: Optional[datetime] = None
    updated_at: Optional[datetime] = None
    processed_at: Optional[datetime] = None
    status: ProcessingStatus = ProcessingStatus.PENDING
    chunk_count: int = 0
    source_hash: Optional[str] = None

@dataclass
class SearchFilters:
    """Filters for knowledge base search"""
    document_types: Optional[List[str]] = None
    tags: Optional[List[str]] = None
    date_from: Optional[datetime] = None
    date_to: Optional[datetime] = None
    file_types: Optional[List[str]] = None
    has_chunks: Optional[bool] = None

class KnowledgeBaseService:
    """
    Knowledge base service for document management and semantic search
    """
    
    def __init__(self, storage_path: Optional[str] = None, vector_service: Optional[VectorSearchService] = None):
        """
        Initialize knowledge base service
        
        Args:
            storage_path: Path for document storage
            vector_service: Vector search service instance
        """
        self.config = self._load_config(storage_path)
        self.storage_path = Path(self.config["storage_path"])
        self.vector_service = vector_service or VectorSearchService()
        
        # Document metadata storage (in production, use proper database)
        self.metadata_file = self.storage_path / "metadata.json"
        self.documents = {}
        
        # Ensure storage directories exist
        self._create_storage_structure()
        
        # Load existing metadata
        self._load_metadata()
        
        logger.info(f"Knowledge base initialized at: {self.storage_path}")
    
    def _load_config(self, storage_path: Optional[str] = None) -> Dict[str, Any]:
        """Load configuration"""
        return {
            "storage_path": storage_path or os.getenv("CATALYST_KB_STORAGE_PATH", "./data/knowledge_base"),
            "max_file_size": int(os.getenv("CATALYST_KB_MAX_FILE_SIZE", str(50 * 1024 * 1024))),  # 50MB
            "supported_types": [
                "text/plain", "text/markdown", "text/html",
                "application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "image/jpeg", "image/png", "image/gif", "image/bmp", "image/tiff"
            ],
            "auto_categorize": os.getenv("CATALYST_KB_AUTO_CATEGORIZE", "true").lower() == "true",
            "auto_tag": os.getenv("CATALYST_KB_AUTO_TAG", "true").lower() == "true",
            "ocr_enabled": OCR_AVAILABLE and os.getenv("CATALYST_KB_OCR_ENABLED", "true").lower() == "true"
        }
    
    def _create_storage_structure(self):
        """Create storage directory structure"""
        self.storage_path.mkdir(parents=True, exist_ok=True)
        
        # Create subdirectories for organization
        (self.storage_path / "documents").mkdir(exist_ok=True)
        (self.storage_path / "processed").mkdir(exist_ok=True)
        (self.storage_path / "temp").mkdir(exist_ok=True)
    
    def _load_metadata(self):
        """Load document metadata from storage"""
        try:
            if self.metadata_file.exists():
                with open(self.metadata_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    # Convert to KnowledgeDocument objects
                    for doc_id, doc_data in data.items():
                        # Convert datetime strings back to datetime objects
                        for field in ['created_at', 'updated_at', 'processed_at']:
                            if doc_data.get(field):
                                doc_data[field] = datetime.fromisoformat(doc_data[field])
                        
                        self.documents[doc_id] = KnowledgeDocument(**doc_data)
                        
                logger.info(f"Loaded {len(self.documents)} documents from metadata")
        except Exception as e:
            logger.error(f"Error loading metadata: {e}")
            self.documents = {}
    
    def _save_metadata(self):
        """Save document metadata to storage"""
        try:
            # Convert to serializable format
            data = {}
            for doc_id, doc in self.documents.items():
                doc_dict = asdict(doc)
                # Convert datetime objects to strings
                for field in ['created_at', 'updated_at', 'processed_at']:
                    if doc_dict.get(field):
                        doc_dict[field] = doc_dict[field].isoformat()
                data[doc_id] = doc_dict
            
            with open(self.metadata_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
                
        except Exception as e:
            logger.error(f"Error saving metadata: {e}")
    
    async def initialize(self) -> bool:
        """
        Initialize the knowledge base service
        
        Returns:
            bool: True if initialization successful
        """
        try:
            # Initialize vector service
            if not await self.vector_service.initialize():
                logger.error("Failed to initialize vector service")
                return False
            
            logger.info("Knowledge base service initialized successfully")
            return True
            
        except Exception as e:
            logger.error(f"Failed to initialize knowledge base service: {e}")
            return False
    
    async def add_document(self, 
                          title: str,
                          content: str = None,
                          file_path: str = None,
                          file_data: bytes = None,
                          document_type: DocumentType = DocumentType.OTHER,
                          tags: List[str] = None,
                          metadata: Dict[str, Any] = None,
                          auto_process: bool = True) -> Optional[str]:
        """
        Add a document to the knowledge base
        
        Args:
            title: Document title
            content: Document content (if not from file)
            file_path: Path to source file
            file_data: Raw file data
            document_type: Document type
            tags: Document tags
            metadata: Additional metadata
            auto_process: Whether to automatically process and index
            
        Returns:
            Optional[str]: Document ID if successful
        """
        try:
            # Generate document ID
            doc_id = self._generate_document_id(title, content or file_path)
            
            # Determine content source
            if content:
                # Direct content
                final_content = content
                file_type = "text/plain"
                file_size = len(content.encode('utf-8'))
                source_hash = hashlib.md5(content.encode('utf-8')).hexdigest()
                stored_file_path = None
                
            elif file_data:
                # File data provided
                result = await self._process_file_data(file_data, title, doc_id)
                if not result:
                    return None
                final_content, file_type, file_size, stored_file_path = result
                source_hash = hashlib.md5(file_data).hexdigest()
                
            elif file_path:
                # File path provided
                result = await self._process_file_path(file_path, doc_id)
                if not result:
                    return None
                final_content, file_type, file_size, stored_file_path = result
                with open(file_path, 'rb') as f:
                    source_hash = hashlib.md5(f.read()).hexdigest()
            else:
                logger.error("No content, file_path, or file_data provided")
                return None
            
            # Auto-categorize if enabled
            if self.config.get("auto_categorize") and document_type == DocumentType.OTHER:
                document_type = await self._auto_categorize(final_content, title)
            
            # Auto-tag if enabled
            if self.config.get("auto_tag") and not tags:
                tags = await self._auto_tag(final_content, title)
            
            # Create document record
            now = datetime.now(timezone.utc)
            document = KnowledgeDocument(
                id=doc_id,
                title=title,
                content=final_content,
                document_type=document_type,
                file_path=stored_file_path,
                file_type=file_type,
                file_size=file_size,
                tags=tags or [],
                metadata=metadata or {},
                created_at=now,
                updated_at=now,
                status=ProcessingStatus.PENDING,
                source_hash=source_hash
            )
            
            # Store document
            self.documents[doc_id] = document
            self._save_metadata()
            
            # Auto-process if enabled
            if auto_process:
                await self.process_document(doc_id)
            
            logger.info(f"Added document: {doc_id} - {title}")
            return doc_id
            
        except Exception as e:
            logger.error(f"Error adding document: {e}")
            return None
    
    def _generate_document_id(self, title: str, content_ref: str = None) -> str:
        """Generate unique document ID"""
        base = f"{title}_{content_ref or ''}_{datetime.now(timezone.utc).isoformat()}"
        return hashlib.md5(base.encode('utf-8')).hexdigest()[:16]
    
    async def _process_file_data(self, file_data: bytes, filename: str, doc_id: str) -> Optional[Tuple[str, str, int, str]]:
        """Process file data and extract content"""
        try:
            # Determine file type
            file_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
            
            # Check file size
            if len(file_data) > self.config["max_file_size"]:
                logger.error(f"File too large: {len(file_data)} bytes")
                return None
            
            # Check supported types
            if file_type not in self.config["supported_types"]:
                logger.error(f"Unsupported file type: {file_type}")
                return None
            
            # Store file
            file_ext = Path(filename).suffix or ".txt"
            stored_path = self.storage_path / "documents" / f"{doc_id}{file_ext}"
            with open(stored_path, 'wb') as f:
                f.write(file_data)
            
            # Extract content
            content = await self._extract_content_from_file(stored_path, file_type)
            
            return content, file_type, len(file_data), str(stored_path)
            
        except Exception as e:
            logger.error(f"Error processing file data: {e}")
            return None
    
    async def _process_file_path(self, file_path: str, doc_id: str) -> Optional[Tuple[str, str, int, str]]:
        """Process file from path and extract content"""
        try:
            source_path = Path(file_path)
            if not source_path.exists():
                logger.error(f"File not found: {file_path}")
                return None
            
            # Check file size
            file_size = source_path.stat().st_size
            if file_size > self.config["max_file_size"]:
                logger.error(f"File too large: {file_size} bytes")
                return None
            
            # Determine file type
            file_type = mimetypes.guess_type(file_path)[0] or "application/octet-stream"
            
            # Check supported types
            if file_type not in self.config["supported_types"]:
                logger.error(f"Unsupported file type: {file_type}")
                return None
            
            # Copy file to storage
            file_ext = source_path.suffix or ".txt"
            stored_path = self.storage_path / "documents" / f"{doc_id}{file_ext}"
            
            with open(source_path, 'rb') as src, open(stored_path, 'wb') as dst:
                dst.write(src.read())
            
            # Extract content
            content = await self._extract_content_from_file(stored_path, file_type)
            
            return content, file_type, file_size, str(stored_path)
            
        except Exception as e:
            logger.error(f"Error processing file path: {e}")
            return None
    
    async def _extract_content_from_file(self, file_path: Path, file_type: str) -> str:
        """Extract text content from file based on type"""
        try:
            if file_type == "text/plain" or file_type == "text/markdown":
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
                    
            elif file_type == "text/html":
                # Basic HTML text extraction (in production, use BeautifulSoup)
                with open(file_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                # Remove HTML tags (basic)
                text = re.sub(r'<[^>]+>', '', html_content)
                return text.strip()
                
            elif file_type == "application/pdf" and PDF_AVAILABLE:
                return await self._extract_pdf_content(file_path)
                
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" and DOCX_AVAILABLE:
                return await self._extract_docx_content(file_path)
                
            elif file_type.startswith("image/") and self.config.get("ocr_enabled"):
                return await self._extract_image_content(file_path)
                
            else:
                logger.warning(f"No content extractor for file type: {file_type}")
                return f"[Content from {file_type} file - processing not supported]"
                
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            return f"[Error extracting content: {str(e)}]"
    
    async def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            content = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        content.append(text)
            
            return '\n\n'.join(content)
        except Exception as e:
            logger.error(f"Error extracting PDF content: {e}")
            return "[Error extracting PDF content]"
    
    async def _extract_docx_content(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            return '\n\n'.join(content)
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {e}")
            return "[Error extracting DOCX content]"
    
    async def _extract_image_content(self, file_path: Path) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting image content: {e}")
            return "[Error extracting image content]"
    
    async def _auto_categorize(self, content: str, title: str) -> DocumentType:
        """Auto-categorize document based on content"""
        # Simple keyword-based categorization
        text = (title + " " + content).lower()
        
        if any(word in text for word in ["guide", "how to", "tutorial", "instruction"]):
            return DocumentType.GUIDANCE
        elif any(word in text for word in ["reference", "documentation", "spec", "standard"]):
            return DocumentType.REFERENCE
        elif any(word in text for word in ["case study", "example", "scenario"]):
            return DocumentType.CASE_STUDY
        elif any(word in text for word in ["template", "form", "format"]):
            return DocumentType.TEMPLATE
        elif any(word in text for word in ["research", "study", "analysis", "findings"]):
            return DocumentType.RESEARCH
        elif any(word in text for word in ["manual", "handbook", "procedures"]):
            return DocumentType.MANUAL
        elif any(word in text for word in ["faq", "questions", "q&a"]):
            return DocumentType.FAQ
        else:
            return DocumentType.OTHER
    
    async def _auto_tag(self, content: str, title: str) -> List[str]:
        """Auto-generate tags for document"""
        # Simple keyword extraction (in production, use NLP)
        text = (title + " " + content).lower()
        
        # Common relationship/therapy keywords
        therapy_keywords = [
            "communication", "relationship", "therapy", "counseling", "conflict",
            "emotional", "intimacy", "trust", "attachment", "boundaries",
            "empathy", "listening", "validation", "support", "growth"
        ]
        
        tags = [keyword for keyword in therapy_keywords if keyword in text]
        
        # Add document type-based tags
        if "guide" in text or "tutorial" in text:
            tags.append("guide")
        if "example" in text or "case" in text:
            tags.append("example")
        if "research" in text or "study" in text:
            tags.append("research")
        
        return list(set(tags))  # Remove duplicates
    
    @handle_kb_errors
    async def process_document(self, document_id: str) -> ProcessingResult:
        """
        Process a document for indexing
        
        Args:
            document_id: Document ID to process
            
        Returns:
            ProcessingResult: Processing results
        """
        start_time = datetime.now(timezone.utc)
        errors = []
        
        try:
            document = self.documents.get(document_id)
            if not document:
                return ProcessingResult(
                    document_id=document_id,
                    success=False,
                    chunks_created=0,
                    processing_time=0,
                    errors=["Document not found"],
                    metadata={}
                )
            
            # Update status
            document.status = ProcessingStatus.PROCESSING
            self._save_metadata()
            
            # Chunk document if necessary
            chunks = self._chunk_document(document.content)
            
            # Prepare metadata for indexing
            base_metadata = {
                "document_id": document_id,
                "title": document.title,
                "document_type": document.document_type.value,
                "tags": document.tags,
                "file_type": document.file_type,
                "created_at": document.created_at.isoformat() if document.created_at else None,
                **document.metadata
            }
            
            # Index in vector database
            success = await self.vector_service.index_document(
                document_id=document_id,
                content=document.content,
                metadata=base_metadata,
                chunk_content=len(chunks) > 1
            )
            
            if success:
                # Update document status
                document.status = ProcessingStatus.INDEXED
                document.processed_at = datetime.now(timezone.utc)
                document.chunk_count = len(chunks)
            else:
                document.status = ProcessingStatus.FAILED
                errors.append("Failed to index in vector database")
            
            self._save_metadata()
            
            processing_time = (datetime.now(timezone.utc) - start_time).total_seconds()
            
            result = ProcessingResult(
                document_id=document_id,
                success=success,
                chunks_created=len(chunks),
                processing_time=processing_time,
                errors=errors,
                metadata=base_metadata
            )
            
            logger.info(f"Processed document {document_id}: {success}, {len(chunks)} chunks, {processing_time:.2f}s")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {document_id}: {e}")
            
            # Update status on error
            if document_id in self.documents:
                self.documents[document_id].status = ProcessingStatus.FAILED
                self._save_metadata()
            
            processing_time = (datetime.now(timezone.utc) - start_time).total_seconds()
            return ProcessingResult(
                document_id=document_id,
                success=False,
                chunks_created=0,
                processing_time=processing_time,
                errors=[str(e)],
                metadata={}
            )
    
    def _chunk_document(self, content: str) -> List[DocumentChunk]:
        """Chunk document content"""
        # Use vector service chunking configuration
        chunk_size = self.vector_service.config.get("chunk_size", 1000)
        
        if len(content) <= chunk_size:
            return [DocumentChunk(
                chunk_id="0",
                content=content,
                chunk_index=0,
                total_chunks=1,
                start_char=0,
                end_char=len(content),
                metadata={}
            )]
        
        chunks = []
        # Simple sentence-boundary chunking
        sentences = re.split(r'[.!?]+', content)
        current_chunk = ""
        chunk_index = 0
        start_char = 0
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                # Save current chunk
                end_char = start_char + len(current_chunk)
                chunks.append(DocumentChunk(
                    chunk_id=str(chunk_index),
                    content=current_chunk.strip(),
                    chunk_index=chunk_index,
                    total_chunks=0,  # Will be updated later
                    start_char=start_char,
                    end_char=end_char,
                    metadata={}
                ))
                
                start_char = end_char
                current_chunk = sentence
                chunk_index += 1
            else:
                current_chunk += sentence
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append(DocumentChunk(
                chunk_id=str(chunk_index),
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                total_chunks=0,
                start_char=start_char,
                end_char=start_char + len(current_chunk),
                metadata={}
            ))
        
        # Update total_chunks for all chunks
        for chunk in chunks:
            chunk.total_chunks = len(chunks)
        
        return chunks
    
    async def search_knowledge_base(self, 
                                  query: str,
                                  filters: SearchFilters = None,
                                  limit: int = 10,
                                  min_similarity: float = 0.3) -> List[SearchResult]:
        """
        Search the knowledge base
        
        Args:
            query: Search query
            filters: Search filters
            limit: Maximum results
            min_similarity: Minimum similarity threshold
            
        Returns:
            List[SearchResult]: Search results
        """
        try:
            # Build vector database filters
            vector_filters = {}
            
            if filters:
                if filters.document_types:
                    vector_filters["document_type"] = {"$in": filters.document_types}
                
                if filters.tags:
                    # For now, simple tag matching (in production, use proper array matching)
                    vector_filters["tags"] = {"$in": filters.tags}
            
            # Search vector database
            results = await self.vector_service.search(
                query=query,
                limit=limit,
                min_similarity=min_similarity,
                filters=vector_filters
            )
            
            # Enhance results with document metadata
            enhanced_results = []
            for result in results:
                doc_id = result.metadata.get("document_id", result.document_id)
                if doc_id in self.documents:
                    doc = self.documents[doc_id]
                    result.metadata.update({
                        "document_title": doc.title,
                        "document_status": doc.status.value,
                        "created_at": doc.created_at.isoformat() if doc.created_at else None
                    })
                
                enhanced_results.append(result)
            
            return enhanced_results
            
        except Exception as e:
            logger.error(f"Error searching knowledge base: {e}")
            return []
    
    async def get_document(self, document_id: str) -> Optional[KnowledgeDocument]:
        """Get document by ID"""
        return self.documents.get(document_id)
    
    async def list_documents(self, 
                           filters: SearchFilters = None,
                           limit: int = 100,
                           offset: int = 0) -> List[KnowledgeDocument]:
        """List documents with optional filters"""
        try:
            documents = list(self.documents.values())
            
            # Apply filters
            if filters:
                if filters.document_types:
                    documents = [d for d in documents if d.document_type.value in filters.document_types]
                
                if filters.tags:
                    documents = [d for d in documents if any(tag in d.tags for tag in filters.tags)]
                
                if filters.date_from:
                    documents = [d for d in documents if d.created_at and d.created_at >= filters.date_from]
                
                if filters.date_to:
                    documents = [d for d in documents if d.created_at and d.created_at <= filters.date_to]
                
                if filters.file_types:
                    documents = [d for d in documents if d.file_type in filters.file_types]
                
                if filters.has_chunks is not None:
                    if filters.has_chunks:
                        documents = [d for d in documents if d.chunk_count > 1]
                    else:
                        documents = [d for d in documents if d.chunk_count <= 1]
            
            # Sort by created_at (newest first)
            documents.sort(key=lambda x: x.created_at or datetime.min, reverse=True)
            
            # Apply pagination
            return documents[offset:offset + limit]
            
        except Exception as e:
            logger.error(f"Error listing documents: {e}")
            return []
    
    async def update_document(self, 
                            document_id: str,
                            title: str = None,
                            content: str = None,
                            document_type: DocumentType = None,
                            tags: List[str] = None,
                            metadata: Dict[str, Any] = None) -> bool:
        """Update document metadata and re-index if necessary"""
        try:
            document = self.documents.get(document_id)
            if not document:
                return False
            
            # Track if content changed (requiring re-indexing)
            content_changed = False
            
            # Update fields
            if title is not None:
                document.title = title
            
            if content is not None:
                if content != document.content:
                    document.content = content
                    content_changed = True
            
            if document_type is not None:
                document.document_type = document_type
            
            if tags is not None:
                document.tags = tags
            
            if metadata is not None:
                document.metadata.update(metadata)
            
            document.updated_at = datetime.now(timezone.utc)
            
            # Re-index if content changed
            if content_changed:
                await self.process_document(document_id)
            
            self._save_metadata()
            return True
            
        except Exception as e:
            logger.error(f"Error updating document {document_id}: {e}")
            return False
    
    async def delete_document(self, document_id: str) -> bool:
        """Delete document and remove from vector database"""
        try:
            document = self.documents.get(document_id)
            if not document:
                return False
            
            # Delete from vector database
            await self.vector_service.delete_document(document_id)
            
            # Delete stored file if exists
            if document.file_path and Path(document.file_path).exists():
                Path(document.file_path).unlink()
            
            # Remove from metadata
            del self.documents[document_id]
            self._save_metadata()
            
            logger.info(f"Deleted document: {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {e}")
            return False
    
    async def get_statistics(self) -> Dict[str, Any]:
        """Get knowledge base statistics"""
        try:
            docs = list(self.documents.values())
            
            # Basic counts
            total_docs = len(docs)
            by_status = {}
            by_type = {}
            by_file_type = {}
            
            total_size = 0
            total_chunks = 0
            
            for doc in docs:
                # Status distribution (handle both enum and string)
                status = doc.status.value if hasattr(doc.status, 'value') else str(doc.status)
                by_status[status] = by_status.get(status, 0) + 1
                
                # Type distribution (handle both enum and string)
                doc_type = doc.document_type.value if hasattr(doc.document_type, 'value') else str(doc.document_type)
                by_type[doc_type] = by_type.get(doc_type, 0) + 1
                
                # File type distribution
                if doc.file_type:
                    by_file_type[doc.file_type] = by_file_type.get(doc.file_type, 0) + 1
                
                # Size and chunks
                if doc.file_size:
                    total_size += doc.file_size
                
                total_chunks += doc.chunk_count
            
            # Vector database stats
            vector_count = await self.vector_service.get_document_count()
            
            return {
                "total_documents": total_docs,
                "total_size_bytes": total_size,
                "total_chunks": total_chunks,
                "vector_count": vector_count,
                "distribution": {
                    "by_status": by_status,
                    "by_type": by_type,
                    "by_file_type": by_file_type
                },
                "avg_chunks_per_doc": total_chunks / max(total_docs, 1),
                "indexed_percentage": (by_status.get("indexed", 0) / max(total_docs, 1)) * 100
            }
            
        except Exception as e:
            logger.error(f"Error getting statistics: {e}")
            return {}
    
    async def health_check(self) -> Dict[str, Any]:
        """Check service health"""
        status = {
            "service": "knowledge_base",
            "status": "healthy",
            "storage_path": str(self.storage_path),
            "document_count": len(self.documents),
            "vector_service": {},
            "errors": []
        }
        
        try:
            # Check storage path
            if not self.storage_path.exists():
                status["errors"].append("Storage path does not exist")
            
            # Check vector service
            status["vector_service"] = await self.vector_service.health_check()
            if status["vector_service"]["status"] != "healthy":
                status["errors"].append("Vector service unhealthy")
            
            # Check if we can process documents
            test_doc_id = await self.add_document(
                title="Health Check Test",
                content="This is a test document for health checking.",
                auto_process=True
            )
            
            if test_doc_id:
                # Clean up test document
                await self.delete_document(test_doc_id)
            else:
                status["errors"].append("Failed to add test document")
            
        except Exception as e:
            status["status"] = "unhealthy"
            status["errors"].append(str(e))
        
        if status["errors"]:
            status["status"] = "unhealthy"
        
        return status

    async def enrich_search_results(self, results: List[SearchResult]) -> List[SearchResult]:
        """
        Enrich search results with document metadata
        
        Args:
            results: List of search results from vector search
            
        Returns:
            List[SearchResult]: Enriched search results
        """
        enriched_results = []
        
        for result in results:
            try:
                # Get document metadata
                document = await self.get_document(result.document_id)
                
                if document:
                    # Enhance the metadata with document information
                    enriched_metadata = {
                        **result.metadata,
                        "title": document.get("title"),
                        "filename": document.get("filename"),
                        "content_type": document.get("content_type"),
                        "document_type": document.get("document_type"),
                        "category": document.get("category"),
                        "tags": document.get("tags", []),
                        "created_at": document.get("created_at"),
                    }
                    
                    # Create enriched result
                    enriched_result = SearchResult(
                        document_id=result.document_id,
                        content=result.content,
                        metadata=enriched_metadata,
                        similarity_score=result.similarity_score,
                        chunk_index=result.chunk_index,
                        total_chunks=result.total_chunks
                    )
                    
                    enriched_results.append(enriched_result)
                else:
                    # If document not found, keep original result
                    enriched_results.append(result)
                    
            except Exception as e:
                logger.error(f"Error enriching search result for document {result.document_id}: {str(e)}")
                # Keep original result in case of error
                enriched_results.append(result)
        
        return enriched_results

def handle_kb_errors(func):
    """
    Decorator for handling knowledge base service errors
    
    Args:
        func: The async function to wrap
        
    Returns:
        Wrapped function with error handling
    """
    @functools.wraps(func)
    async def wrapper(*args, **kwargs):
        try:
            return await func(*args, **kwargs)
        except Exception as e:
            # Get the function name for better error context
            func_name = func.__name__
            
            # Log the error with detailed context
            logger.error(f"Knowledge base error in {func_name}: {str(e)}", 
                         extra={
                             "args": args[1:],  # Skip self
                             "kwargs": kwargs,
                             "error_type": type(e).__name__
                         },
                         exc_info=True)
            
            # Return a proper error response instead of raising
            if func_name.startswith("search"):
                return []
            elif func_name.startswith("get_"):
                return None
            elif func_name.startswith("index_") or func_name.startswith("process_"):
                return {
                    "success": False,
                    "error": str(e),
                    "error_type": type(e).__name__
                }
            else:
                # For update/delete operations
                return {
                    "success": False,
                    "error": str(e)
                }
    
    return wrapper

# Convenience function
def create_knowledge_base_service(**kwargs) -> KnowledgeBaseService:
    """Create and initialize knowledge base service"""
    return KnowledgeBaseService(**kwargs)
